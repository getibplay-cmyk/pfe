from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ARTIFACTS = ROOT / "artifacts"
OUTPUT = ROOT / "report" / "Compte_rendu_TP1_TP2_Systemes_de_recommandation.docx"

TP1 = json.loads((ARTIFACTS / "metrics_tp1.json").read_text(encoding="utf-8"))
TP2 = json.loads((ARTIFACTS / "metrics_tp2.json").read_text(encoding="utf-8"))


# Preset: standard_business_brief.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "65727E"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "7A5A00"
GREEN = "1F5D42"
RED = "9B1C1C"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, *, size: float | None = None, bold: bool | None = None,
                 italic: bool | None = None, color: str | None = None,
                 name: str = "Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, bottom: int = 80,
                     start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "D6DADE", size: str = "5") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def add_field(paragraph, instruction: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.0

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = rgb(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table_cite = styles.add_style("Table Citation", 1)
    table_cite.base_style = normal
    table_cite.font.name = "Calibri"
    table_cite._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    table_cite._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    table_cite.font.size = Pt(8.5)
    table_cite.font.italic = True
    table_cite.font.color.rgb = rgb(MUTED)
    table_cite.paragraph_format.space_before = Pt(4)
    table_cite.paragraph_format.space_after = Pt(4)

    formula = styles.add_style("Formula", 1)
    formula.base_style = normal
    formula.font.name = "Cambria Math"
    formula._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    formula._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    formula.font.size = Pt(10.5)
    formula.font.color.rgb = rgb(INK)
    formula.paragraph_format.left_indent = Inches(0.35)
    formula.paragraph_format.space_before = Pt(4)
    formula.paragraph_format.space_after = Pt(8)

    code = styles.add_style("Code", 1)
    code.base_style = normal
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code.font.size = Pt(9)
    code.font.color.rgb = rgb(INK)
    code.paragraph_format.left_indent = Inches(0.25)
    code.paragraph_format.right_indent = Inches(0.25)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(7)
    code.paragraph_format.line_spacing = 1.0


def configure_numbering(doc: Document) -> tuple[int, int, int]:
    numbering = doc.part.numbering_part.element

    def make_abstract(abstract_id: int, num_fmt: str, text: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        fmt = OxmlElement("w:numFmt")
        fmt.set(qn("w:val"), num_fmt)
        lvl.append(fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text)
        lvl.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        lvl.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "720")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "720")
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "160")
        spacing.set(qn("w:line"), "280")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
        # OOXML requires every abstractNum before the concrete num elements.
        first_num = numbering.find(qn("w:num"))
        if first_num is None:
            numbering.append(abstract)
        else:
            numbering.insert(list(numbering).index(first_num), abstract)

    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    next_abs = max(existing_abs, default=0) + 1
    bullet_abs, decimal_abs = next_abs, next_abs + 1
    make_abstract(bullet_abs, "bullet", "•")
    make_abstract(decimal_abs, "decimal", "%1.")

    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    next_num = max(existing_num, default=0) + 1

    def make_num(num_id: int, abstract_id: int, *, restart: bool = False) -> None:
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        if restart:
            override = OxmlElement("w:lvlOverride")
            override.set(qn("w:ilvl"), "0")
            start_override = OxmlElement("w:startOverride")
            start_override.set(qn("w:val"), "1")
            override.append(start_override)
            num.append(override)
        numbering.append(num)

    make_num(next_num, bullet_abs)
    make_num(next_num + 1, decimal_abs)
    # A separate instance restarts the independent execution procedure at 1.
    make_num(next_num + 2, decimal_abs, restart=True)
    return next_num, next_num + 1, next_num + 2


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_ref = OxmlElement("w:numId")
    num_ref.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_ref])


def add_bullet(doc: Document, text: str, *, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.add_run(text)


def add_number(doc: Document, text: str, *, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_num(p, num_id)
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, *, accent: str = BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH_DXA])
    set_table_borders(table, color="D8DEE5", size="4")
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=accent)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.08


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int],
              *, font_size: float = 9.3, header_fill: str = LIGHT) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run_font(r, size=font_size, bold=True, color=INK)
    for row_values in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_values):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            r = p.add_run(str(value))
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph("Source : calculs reproductibles des notebooks fournis.", style="Table Citation")


def add_figure(doc: Document, image_path: Path, caption: str, *, width: float = 6.15) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    doc.add_paragraph(caption, style="Caption")


def add_page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def format_rmse(value: float) -> str:
    return f"{value:.4f}"


def format_ndcg(value: float) -> str:
    return f"{value:.4f}"


def build_report() -> Path:
    doc = Document()
    configure_styles(doc)
    bullet_num, decimal_num, execution_num = configure_numbering(doc)

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = False

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("SYSTÈMES DE RECOMMANDATION  |  TP 1 & TP 2")
    set_run_font(hr, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    fr = fp.add_run("Compte rendu  •  ")
    set_run_font(fr, size=8.5, color=MUTED)
    add_field(fp, "PAGE")

    # First-page memo masthead.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("COMPTE RENDU TECHNIQUE")
    set_run_font(r, size=23, bold=True, color="000000")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("TP 1 — Systèmes de recommandation\nTP 2 — Filtrage collaboratif")
    set_run_font(r, size=15, bold=True, color=DARK_BLUE)

    metadata = [
        ("Formation", "Master Sciences des Données"),
        ("Année", "2025–2026"),
        ("Jeux de données", "MovieLens Latest Small et MovieLens 100K"),
        ("Environnement", "Python 3.12 • NumPy 2.3.5 • pandas 2.2.3 • scikit-learn 1.8"),
        ("Date du rendu", "1er septembre 2026"),
    ]
    for label, value in metadata:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        lr = p.add_run(f"{label} : ")
        set_run_font(lr, size=10.5, bold=True)
        vr = p.add_run(value)
        set_run_font(vr, size=10.5)

    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(10)
    rule.paragraph_format.space_after = Pt(12)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    doc.add_heading("Résumé exécutif", level=1)
    add_callout(
        doc,
        "Résultat principal",
        "Le centrage des notes est le levier le plus déterminant. Sur le TP 1, la SVD centrée de rang 5 "
        "obtient la meilleure RMSE (0,9162). Sur le TP 2, le filtrage item-based centré avec k = 200 "
        "atteint la meilleure MSE de test (0,9416).",
        accent=GREEN,
    )
    doc.add_paragraph(
        "Les deux sujets ont été implémentés dans deux notebooks Colab autonomes. Les données sont téléchargées "
        "à l’exécution depuis la source officielle GroupLens ; elles ne sont pas incluses dans les livrables. "
        "Les expériences utilisent une graine fixe (42), des divisions apprentissage/test sans fuite et des "
        "tests unitaires sur les opérations sensibles."
    )
    add_bullet(doc, "TP 1 : 100 836 notes, 610 utilisateurs, 9 724 films ; séparation globale 80/20.", num_id=bullet_num)
    add_bullet(doc, "TP 2 : 100 000 notes, 943 utilisateurs, 1 682 films ; 10 notes de test par utilisateur.", num_id=bullet_num)
    add_bullet(doc, "Évaluation : RMSE/MSE pour la précision des notes et NDCG@10 pour la qualité du classement.", num_id=bullet_num)

    add_page_break(doc)
    doc.add_heading("1. Cadre expérimental", level=1)
    doc.add_heading("1.1 Objectifs couverts", level=2)
    doc.add_paragraph(
        "Le TP 1 demande de construire une chaîne complète allant de la matrice d’utilité au filtrage collaboratif "
        "user-based et item-based, puis à la factorisation SVD. Le TP 2 approfondit la même famille de méthodes : "
        "split par utilisateur, similarité cosinus, prédictions pondérées, influence de k, débiaisage et recherche "
        "de films voisins selon cosinus et Pearson."
    )
    for text in (
        "construire les matrices utilisateur–film sans confondre une note nulle avec une valeur manquante ;",
        "calculer les similarités et exclure l’entité elle-même des voisinages ;",
        "prédire uniquement à partir des notes disponibles dans l’apprentissage ;",
        "comparer les modèles sur des observations jamais utilisées pour l’ajustement ;",
        "interpréter les écarts entre erreur de note et qualité de classement.",
    ):
        add_bullet(doc, text, num_id=bullet_num)

    doc.add_heading("1.2 Reproductibilité et contrôle de qualité", level=2)
    add_table(
        doc,
        ["Contrôle", "Mise en œuvre", "Résultat"],
        [
            ["Données", "GroupLens puis miroir automatique en secours", "Téléchargement robuste dans Colab"],
            ["Aléatoire", "Graine SEED = 42", "Splits reproductibles"],
            ["Valeurs manquantes", "Zéros exclus des moyennes", "Biais non dilué par la sparsité"],
            ["Fuite de données", "Masques train/test disjoints", "Aucune observation partagée"],
            ["Tests", "5 tests unitaires + intégration", "5/5 réussis"],
            ["Notebooks", "Exécution complète avec sorties embarquées", "Aucune cellule en erreur"],
        ],
        [1700, 4200, 3460],
        font_size=9.1,
    )
    doc.add_paragraph(
        "Les fonctions communes sont regroupées dans recommender_core.py et également embarquées dans chaque "
        "notebook afin qu’un fichier .ipynb puisse être ouvert et exécuté seul dans Google Colab."
    )

    doc.add_heading("1.3 Mesures", level=2)
    doc.add_paragraph(
        "La RMSE pénalise davantage les grandes erreurs et reste exprimée dans l’unité de la note. La MSE est son "
        "carré. NDCG@10 compare l’ordre des dix recommandations au classement idéal et mesure donc une propriété "
        "différente de la calibration numérique."
    )
    doc.add_paragraph("RMSE = √[(1 / |T|) · Σ(u,i)∈T (rᵤᵢ − r̂ᵤᵢ)²]", style="Formula")
    doc.add_paragraph("NDCG@10 = DCG@10 / IDCG@10", style="Formula")

    add_page_break(doc)
    doc.add_heading("2. TP 1 — Systèmes de recommandation", level=1)
    doc.add_heading("2.1 Données et matrice d’utilité (questions 1.a–1.b)", level=2)
    doc.add_paragraph(
        "MovieLens Latest Small fournit 100 836 évaluations explicites. Après le split global 80/20, la matrice "
        "d’apprentissage a une densité de seulement 1,360 %, ce qui confirme que l’essentiel de la matrice est "
        "inconnu. Les identifiants d’utilisateurs et de films sont convertis en indices contigus, tout en conservant "
        "les mappings inverses pour restituer les titres."
    )
    add_table(
        doc,
        ["Indicateur", "Valeur", "Lecture"],
        [
            ["Évaluations", f"{TP1['dataset']['ratings']:,}".replace(",", " "), "Toutes les notes disponibles"],
            ["Utilisateurs", str(TP1["dataset"]["users"]), "Lignes de la matrice"],
            ["Films", f"{TP1['dataset']['items']:,}".replace(",", " "), "Colonnes de la matrice"],
            ["Densité train", f"{TP1['dataset']['train_density_pct']:.3f} %", "Matrice très creuse"],
        ],
        [2400, 1800, 5160],
    )

    doc.add_heading("2.2 Similarité et prédiction (questions 1.c–1.f)", level=2)
    doc.add_paragraph(
        "Les produits scalaires utilisateur–utilisateur et film–film sont calculés par multiplication matricielle. "
        "Les normes pré-calculées transforment ensuite ces produits en cosinus sans recalculer les vecteurs. "
        "Pour une paire cible, les voisins sont triés, l’auto-similarité est exclue et seuls les voisins ayant une "
        "note observée contribuent à la prédiction."
    )
    doc.add_paragraph("sim(a,b) = (a · b) / (‖a‖₂ · ‖b‖₂)", style="Formula")
    doc.add_paragraph("r̂ᵤᵢ = Σᵥ sim(u,v) · rᵥᵢ / Σᵥ |sim(u,v)|", style="Formula")
    add_callout(
        doc,
        "Cas sans voisin exploitable",
        "Lorsque le dénominateur est nul, la prédiction se replie sur la moyenne globale. Les scores sont ensuite "
        "bornés dans l’intervalle valide [0,5 ; 5].",
        accent=GOLD,
    )

    doc.add_heading("2.3 Résultats du filtrage collaboratif (questions 1.g–1.j)", level=2)
    tp1_cf_rows = [
        ["User-based brut, k=10", format_rmse(TP1["raw_cf_rmse"]["User-based k=10"]), format_ndcg(TP1["raw_cf_ndcg10"]["User-based k=10"])],
        ["Item-based brut, k=10", format_rmse(TP1["raw_cf_rmse"]["Item-based k=10"]), format_ndcg(TP1["raw_cf_ndcg10"]["Item-based k=10"])],
        ["User-based centré, k=10", format_rmse(TP1["centered_cf_rmse"]["User-based centré k=10"]), format_ndcg(TP1["centered_cf_ndcg10"]["User-based centré k=10"])],
        ["Item-based centré, k=10", format_rmse(TP1["centered_cf_rmse"]["Item-based centré k=10"]), format_ndcg(TP1["centered_cf_ndcg10"]["Item-based centré k=10"])],
    ]
    add_table(doc, ["Modèle", "RMSE", "NDCG@10"], tp1_cf_rows, [5200, 2080, 2080])
    doc.add_paragraph(
        "Le centrage réduit la RMSE user-based de 0,1046 et la RMSE item-based de 0,0539. La variante item-based "
        "centrée est la meilleure des quatre (0,9435). Les NDCG@10 des méthodes locales restent faibles : avec un "
        "split global et une matrice très creuse, une grande partie des items de test ne reçoit que peu de signal de voisinage."
    )

    doc.add_heading("2.4 Factorisation SVD (questions 1.k–1.m)", level=2)
    doc.add_paragraph(
        "La matrice d’apprentissage est décomposée par SVD puis reconstruite à différents rangs N. Deux variantes "
        "sont évaluées : la matrice brute où les zéros sont factorisés tels quels, et la matrice centrée par la moyenne "
        "de chaque utilisateur, où seuls les résidus observés sont conservés."
    )
    doc.add_paragraph("R ≈ Uₙ · diag(σ₁,…,σₙ) · Vₙᵀ", style="Formula")
    add_table(
        doc,
        ["Variante", "Rang optimal", "RMSE", "NDCG@10"],
        [
            ["SVD brute", str(TP1["svd_raw"]["best_rank"]), format_rmse(TP1["svd_raw"]["rmse"]), format_ndcg(TP1["svd_raw"]["ndcg10"])],
            ["SVD centrée", str(TP1["svd_centered"]["best_rank"]), format_rmse(TP1["svd_centered"]["rmse"]), format_ndcg(TP1["svd_centered"]["ndcg10"])],
        ],
        [3000, 1800, 2280, 2280],
    )
    add_figure(
        doc,
        ARTIFACTS / "tp1_comparaison_svd.png",
        "Figure 1 — Influence du rang et distribution des scores : SVD brute contre SVD centrée.",
    )
    doc.add_paragraph(
        "Le meilleur rang brut est 10, tandis que le meilleur rang après centrage est 5. Au-delà, la RMSE augmente "
        "progressivement : les composantes supplémentaires commencent à reconstruire du bruit et les zéros structurels. "
        "La SVD centrée obtient la meilleure RMSE de tout le TP (0,9162)."
    )

    add_page_break(doc)
    doc.add_heading("2.5 Distribution des scores et analyse (questions 1.n–1.o)", level=2)
    add_figure(
        doc,
        ARTIFACTS / "tp1_distribution_svd_brute.png",
        "Figure 2 — Distribution des scores SVD brute sur les observations de test.",
        width=5.9,
    )
    doc.add_paragraph(
        "La SVD brute traite les zéros de la matrice comme de faibles notes. Elle produit donc des estimations "
        "fortement comprimées vers le bas, souvent hors de l’échelle utile : la calibration est mauvaise et la RMSE "
        "atteint 2,7874. Malgré cela, son NDCG@10 (0,2433) est le meilleur des modèles testés. Cette différence n’est "
        "pas contradictoire : les facteurs latents peuvent encore ordonner des films populaires de façon pertinente "
        "tout en prédisant de mauvaises valeurs absolues."
    )
    add_callout(
        doc,
        "Interprétation des métriques",
        "La RMSE mesure « combien » la note est fausse ; NDCG mesure « dans quel ordre » les films sont proposés. "
        "Une méthode peut donc être mal calibrée mais relativement efficace pour classer.",
        accent=BLUE,
    )
    doc.add_heading("2.6 Réponse synthétique au TP 1", level=2)
    for text in (
        "Le cosinus permet un calcul cohérent pour les deux orientations user-based et item-based.",
        "À k=10, l’item-based centré domine les variantes de filtrage collaboratif en RMSE.",
        "La SVD brute doit être évitée pour la prédiction de notes sur une matrice où zéro signifie « absent ».",
        "Le centrage par utilisateur corrige le biais d’échelle et donne la meilleure RMSE globale avec une SVD de rang 5.",
        "Le choix du modèle final dépend de l’objectif : précision numérique ou classement des recommandations.",
    ):
        add_number(doc, text, num_id=decimal_num)

    add_page_break(doc)
    doc.add_heading("3. TP 2 — Filtrage collaboratif", level=1)
    doc.add_heading("3.1 Données, sparsité et split par utilisateur", level=2)
    doc.add_paragraph(
        "MovieLens 100K contient exactement 100 000 évaluations, fournies par 943 utilisateurs sur 1 682 films. "
        "La matrice a une densité de 6,3047 %, soit une sparsité stricte de 93,6953 %. Pour chaque utilisateur, dix "
        "notes sont déplacées vers le jeu de test. Ce protocole garantit une évaluation personnalisée pour tous les "
        "utilisateurs suffisamment actifs."
    )
    add_table(
        doc,
        ["Indicateur", "Valeur", "Conséquence"],
        [
            ["Évaluations", f"{TP2['dataset']['ratings']:,}".replace(",", " "), "Signal explicite"],
            ["Utilisateurs", str(TP2["dataset"]["users"]), "943 profils évalués"],
            ["Films", f"{TP2['dataset']['items']:,}".replace(",", " "), "Catalogue modéré"],
            ["Densité", f"{TP2['dataset']['density_pct']:.4f} %", "Voisinages partiels"],
            ["Sparsité", f"{100-TP2['dataset']['density_pct']:.4f} %", "Les valeurs manquantes dominent"],
        ],
        [2300, 1900, 5160],
    )

    doc.add_heading("3.2 Similarités et premier voisin", level=2)
    doc.add_paragraph(
        f"Les similarités cosinus sont calculées sur la matrice d’apprentissage. Le plus proche voisin de l’utilisateur "
        f"1 est l’utilisateur {TP2['nearest_user_to_user_1']['user_id']}, avec un cosinus de "
        f"{TP2['nearest_user_to_user_1']['cosine']:.4f}. L’auto-similarité est explicitement neutralisée avant le tri."
    )
    doc.add_heading("3.3 Prédictions simples", level=2)
    simple_rows = [[row["modèle"], f"{row['MSE']:.4f}", f"{row['RMSE']:.4f}"] for row in TP2["simple"]]
    add_table(doc, ["Modèle", "MSE test", "RMSE test"], simple_rows, [5200, 2080, 2080])
    doc.add_paragraph(
        "Ces premières versions agrègent tous les voisins sans centrage. Elles servent de référence, mais restent "
        "sensibles aux différences de sévérité entre utilisateurs et aux niveaux moyens des films."
    )

    add_page_break(doc)
    doc.add_heading("3.4 Influence du nombre de voisins k", level=2)
    raw_by_k: dict[int, dict[str, float]] = {}
    for row in TP2["raw_curve_test"]:
        raw_by_k.setdefault(int(row["k"]), {})[row["modèle"]] = row["MSE"]
    rows = []
    for k in sorted(raw_by_k):
        rows.append([str(k), f"{raw_by_k[k]['User-based']:.4f}", f"{raw_by_k[k]['Item-based']:.4f}"])
    add_table(doc, ["k", "MSE user-based", "MSE item-based"], rows, [1800, 3780, 3780])
    add_figure(
        doc,
        ARTIFACTS / "tp2_influence_k_brut.png",
        "Figure 3 — Influence de k sur l’erreur d’apprentissage et de test des modèles bruts.",
        width=5.95,
    )
    doc.add_paragraph(
        "L’erreur test décroît nettement entre k=5 et k=50 puis s’aplatit. Dans la plage étudiée, le meilleur modèle "
        "brut est item-based avec k=200 (MSE 1,0615). La courbe ne présente pas de remontée avant 200 : les voisinages "
        "larges stabilisent ici les estimations et compensent la sparsité."
    )

    add_page_break(doc)
    doc.add_heading("3.5 Débiaisage : centrage + top-k", level=2)
    doc.add_paragraph(
        "Le modèle user-based soustrait à chaque utilisateur sa moyenne observée, puis la réinjecte après l’agrégation. "
        "La version item-based applique le même principe avec la moyenne observée de chaque film. Les zéros restent "
        "des valeurs manquantes et ne participent jamais au calcul des moyennes."
    )
    centered_by_k: dict[int, dict[str, float]] = {}
    for row in TP2["centered_curve_test"]:
        centered_by_k.setdefault(int(row["k"]), {})[row["modèle"]] = row["MSE"]
    rows = []
    for k in sorted(centered_by_k):
        rows.append([
            str(k),
            f"{centered_by_k[k]['User-based centré']:.4f}",
            f"{centered_by_k[k]['Item-based centré']:.4f}",
        ])
    add_table(doc, ["k", "MSE user centré", "MSE item centré"], rows, [1800, 3780, 3780])
    add_figure(
        doc,
        ARTIFACTS / "tp2_comparaison_k_debiaisage.png",
        "Figure 4 — Erreur test : modèles bruts et modèles centrés selon k.",
        width=5.95,
    )
    doc.add_paragraph(
        "Le débiaisage améliore systématiquement les résultats. À k=200, la MSE item-based passe de 1,0615 à 0,9416, "
        "soit une baisse relative de 11,3 %. La variante user-based centrée atteint 0,9481, très proche du meilleur score."
    )

    add_page_break(doc)
    doc.add_heading("3.6 Films similaires : cosinus et Pearson", level=2)
    doc.add_paragraph(
        "Le cosinus mesure la proximité des profils de notes, mais reste influencé par le niveau moyen et la popularité. "
        "Pearson centre implicitement les profils sur les utilisateurs communs ; il met davantage en avant les écarts "
        "concordants. Les deux mesures produisent donc des listes voisines, sans être identiques."
    )
    neighbor_rows = []
    for title in ("Toy Story", "GoldenEye", "Muppet Treasure Island"):
        neighbor_rows.append([
            title,
            "\n".join(TP2["cosine_neighbors"][title]),
            "\n".join(TP2["pearson_neighbors"][title]),
        ])
    add_table(
        doc,
        ["Film cible", "6 voisins — cosinus", "6 voisins — Pearson"],
        neighbor_rows,
        [1850, 3755, 3755],
        font_size=8.3,
    )
    doc.add_paragraph(
        "Pour GoldenEye, les deux métriques s’accordent fortement sur Under Siege, True Lies, Cliffhanger, Top Gun et "
        "Die Hard: With a Vengeance. Pour Toy Story et Muppet Treasure Island, Pearson introduit davantage de films "
        "proches par tonalité ou public, tandis que le cosinus conserve un signal de co-popularité."
    )
    add_callout(
        doc,
        "Précaution d’interprétation",
        "Une forte similarité ne signifie pas causalité ni équivalence éditoriale. Un seuil minimal de co-évaluations "
        "serait souhaitable avant une mise en production.",
        accent=GOLD,
    )

    add_page_break(doc)
    doc.add_heading("4. Synthèse comparative", level=1)
    synthesis_rows = [
        ["TP 1", "SVD centrée, rang 5", "RMSE", format_rmse(TP1["svd_centered"]["rmse"]), "Meilleure calibration"],
        ["TP 1", "SVD brute, rang 10", "NDCG@10", format_ndcg(TP1["svd_raw"]["ndcg10"]), "Meilleur classement"],
        ["TP 2", "Item-based brut, k=200", "MSE", f"{TP2['best_raw']['mse']:.4f}", "Meilleur sans centrage"],
        ["TP 2", "Item-based centré, k=200", "MSE", f"{TP2['best_centered']['mse']:.4f}", "Meilleur global TP 2"],
    ]
    add_table(
        doc,
        ["Sujet", "Modèle", "Mesure", "Score", "Lecture"],
        synthesis_rows,
        [1100, 2950, 1500, 1450, 2360],
        font_size=8.8,
    )
    doc.add_heading("4.1 Enseignements", level=2)
    for text in (
        "Le traitement correct des valeurs manquantes est aussi important que l’algorithme choisi.",
        "Le centrage corrige les biais de notation et améliore fortement la précision sur les deux jeux.",
        "Le paramètre k arbitre variance et stabilité ; sur MovieLens 100K, les voisinages larges restent bénéfiques jusqu’à 200.",
        "La SVD est efficace après correction de la moyenne, mais la SVD brute confond absence de note et préférence faible.",
        "RMSE/MSE et NDCG répondent à des objectifs distincts et doivent être examinées conjointement.",
    ):
        add_bullet(doc, text, num_id=bullet_num)

    doc.add_heading("4.2 Limites", level=2)
    add_table(
        doc,
        ["Limite", "Effet possible", "Amélioration proposée"],
        [
            ["Un seul split", "Variance du score", "Validation croisée ou plusieurs graines"],
            ["Cosinus non régularisé", "Voisins fragiles", "Shrinkage selon le nombre de co-notes"],
            ["Pas de biais global explicite", "Calibration imparfaite", "Modèle μ + bᵤ + bᵢ + résidu"],
            ["Pas de variables de contenu", "Cold start", "Hybridation genres/tags"],
            ["Pas de contrainte métier", "Top-N peu diversifié", "Diversité, nouveauté et couverture"],
        ],
        [2600, 2700, 4060],
        font_size=8.8,
    )

    doc.add_heading("4.3 Choix recommandé", level=2)
    add_callout(
        doc,
        "Pour la prédiction de notes",
        "Retenir la SVD centrée de rang 5 sur Latest Small et l’item-based centré avec k=200 sur MovieLens 100K, "
        "puis confirmer le choix sur plusieurs splits.",
        accent=GREEN,
    )

    add_page_break(doc)
    doc.add_heading("5. Conclusion", level=1)
    doc.add_paragraph(
        "Les deux TP ont été réalisés intégralement, de l’ingestion des données à l’interprétation des résultats. "
        "Les implémentations montrent que les méthodes de voisinage restent compétitives lorsqu’elles sont correctement "
        "centrées et dotées d’un k suffisant. La factorisation SVD apporte la meilleure précision de note sur TP 1, "
        "mais uniquement après correction du biais utilisateur."
    )
    doc.add_paragraph(
        "Les notebooks livrés sont autonomes, commentés, exécutables dans Google Colab et accompagnés de sorties déjà "
        "calculées. Le module Python commun, les métriques JSON et les tests unitaires permettent de reprendre les "
        "expériences dans un dépôt GitHub sans dépendre du présent document."
    )

    doc.add_heading("6. Livrables et mode d’exécution", level=1)
    add_number(doc, "Ouvrir le notebook TP1_Systemes_de_recommandation.ipynb ou TP2_Filtrage_collaboratif.ipynb dans Colab.", num_id=execution_num)
    add_number(doc, "Choisir « Exécution > Tout exécuter ». Le notebook télécharge le jeu GroupLens requis et bascule automatiquement vers un miroir si la source ne répond pas.", num_id=execution_num)
    add_number(doc, "Conserver la connexion internet le temps du premier téléchargement ; les exécutions suivantes réutilisent /content/data.", num_id=execution_num)
    add_number(doc, "Les métriques et figures sont écrites dans le dossier artifacts du runtime.", num_id=execution_num)
    doc.add_paragraph("Structure livrée :", style="Table Citation")
    doc.add_paragraph(
        "notebooks/TP1_Systemes_de_recommandation.ipynb\n"
        "notebooks/TP2_Filtrage_collaboratif.ipynb\n"
        "src/recommender_core.py\n"
        "tests/test_recommender_core.py\n"
        "artifacts/metrics_tp1.json et metrics_tp2.json\n"
        "report/Compte_rendu_TP1_TP2_Systemes_de_recommandation.pdf",
        style="Code",
    )

    doc.add_heading("7. Références", level=1)
    references = [
        "GroupLens Research. MovieLens Latest Datasets. https://grouplens.org/datasets/movielens/latest/",
        "GroupLens Research. MovieLens 100K Dataset. https://grouplens.org/datasets/movielens/100k/",
        "Sarwar, B. et al. Item-Based Collaborative Filtering Recommendation Algorithms. WWW, 2001.",
        "Koren, Y., Bell, R., Volinsky, C. Matrix Factorization Techniques for Recommender Systems. IEEE Computer, 2009.",
        "Järvelin, K., Kekäläinen, J. Cumulated Gain-Based Evaluation of IR Techniques. ACM TOIS, 2002.",
    ]
    for ref in references:
        add_bullet(doc, ref, num_id=bullet_num)

    # Core properties and update hint for fields.
    doc.core_properties.title = "Compte rendu TP 1 et TP 2 — Systèmes de recommandation"
    doc.core_properties.subject = "Filtrage collaboratif, SVD et MovieLens"
    doc.core_properties.author = "Compte rendu reproductible"
    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_report())
